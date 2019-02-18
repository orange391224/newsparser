# -*- coding: utf-8 -*-
from docxtpl import DocxTemplate
import docx


class RenderDocx(DocxTemplate):


    def __init__(self,file):
        super().__init__(file)


    def render_links(self,context):
        for p in self.docx.paragraphs:
            if p.text[:4] == 'url:':
                font = p.runs[0].font
                title = p.text[4:]
                p.text = ''
                self.add_hyperlink(p, title, context['url_link'],font)
            if p.text[:4] == 'http':
                font = p.runs[0].font
                p.text = ''
                self.add_hyperlink(p, context['url_link'], context['url_link'],font)



    def attach_to_main(self):
        sub_doc = self.docx
        merged_document = docx.Document('main_doc.docx')
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
        self.docx = merged_document


    def add_hyperlink(self,paragraph, text, url,font_style):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.name = font_style.name
        r.font.size = font_style.size
        r.font.underline = font_style.underline
        r.font.color.rgb = font_style.color.rgb
        # r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK

        # r.font.underline = True

        return hyperlink