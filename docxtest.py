# -*- coding: utf-8 -*-
from docxtpl import DocxTemplate
from docx import Document
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it

    #r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    #r.font.underline = True

    return hyperlink


doc = DocxTemplate("my_word_template.docx")

context = { #'section' : 'Финансово-банковский сектор',
            'title' : "url:МИНФИН: ПРОМСВЯЗЬБАНК НЕ БУДЕТ ЕДИНСТВЕННЫМ БАНКОМ ДЛЯ ГОСОБОРОНЗАКАЗА",
            'mass_media' : '(Коммерсант)',
            'date' : '10.07.16',
            'url' : 'http://ya.ru',
            'news_text' : 'Промсвязьбанк не будет единственным банком, который будет обслуживать гособоронзаказ. Об этом на заседании комитета Госдумы по финансовому рынку заявил замминистра финансов РФ Алексей Моисеев. Президент и правительство определили в качестве головного банка Промсвязьбанк, но законопроект не предполагает исключительности этого банка. Есть другие банки, которые находятся под санкциями, и некоторые из них являются уполномоченными по гособоронзаказу. Никто, ни правительство, ни ЦБ не видит необходимости исключения этих банков,– цитирует его'.replace('\xd0',' ')
            }
doc.render(context)
doc.save('demo_render.docx')


def combine_word_documents(input_files):
    """
    :param input_files: an iterable with full paths to docs
    :return: a Document object with the merged files
    """
    for filnr, file in enumerate(input_files):
        # in my case the docx templates are in a FileField of Django, add the MEDIA_ROOT, discard the next 2 lines if not appropriate for you.

        if filnr == 0:
            merged_document = Document(file)
#            merged_document.add_page_break()

        else:
            sub_doc = Document(file)

            # Don't add a page break if you've reached the last file.
            if filnr < len(input_files)-1:
                sub_doc.add_page_break()

            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

    return merged_document


input_files = ['main_doc.docx','demo_render.docx']

d = combine_word_documents(input_files)

for p in d.paragraphs:
    if p.text[:4] == 'url:':
        title = p.text[4:]
        p.text = ''
        add_hyperlink(p,title,context['url'])


d.save('main_doc.docx')